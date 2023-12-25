""" STANDARD + 3RD PARTY PYTHON IMPORTS """

from datetime import datetime

import docx
import pandas as pd
import os
import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches
from docx.shared import Pt


def isfloat(num):
    try:
        float(num)
        return True
    except ValueError:
        return False


def find_number(text):
    num = re.findall(r"[0-9]+", text)
    return " ".join(num)


def read_text_file_signalized(file):
    # read data
    df = pd.read_fwf(file, sep=" ", header=None)

    # merge all columns
    df["combined"] = df.apply(lambda x: "	".join(x.astype(str)), axis=1)

    # drop all other columns
    df = df.drop(columns=df.columns.difference(["combined"]), axis=1)

    # removing nan files from the column
    df["combined"] = df["combined"].str.replace("nan", "0")

    # split into columns
    df = df["combined"].str.split("\t", expand=True)

    # Filtered Rows containing 'Storage Length', 'Lane Group', 'Approach','Storage Length', 'Approach', 'Intersection Signal','Total Delay', 'LOS', 'Queue Length 95'
    df = df[
        (
            df[0].apply(
                lambda val: any(
                    s in str(val)
                    for s in [
                        "Storage Length",
                        "Lane Group",
                        "Approach",
                        "Approach",
                        "Intersection Signal",
                        "Total Delay",
                        "LOS",
                        "Queue Length 95",
                    ]
                )
            )
        )
    ]

    # Transposed df and filter the columns
    df_transposed = df.T
    df_transposed = df_transposed[
        (
            df_transposed[2].apply(
                lambda val: any(s in str(val) for s in ["U", "L", "T", "R"])
            )
        )
    ]
    df = df_transposed.T

    # Promoted row 1 to header in df
    df.columns = df.iloc[0]
    df.drop(df.index[0], inplace=True)

    # Trim whitespaces in columns and remove empty columns
    df.columns = df.columns.str.strip()
    df = df.loc[:, df.columns != ""]

    # Reset Index
    df.reset_index(drop=True, inplace=True)

    # Extract last row
    df.replace(r"^\s*$", None, regex=True, inplace=True)
    df_inter = df[-1:]
    df_inter = df_inter.dropna(axis=1, how="all")
    Intersection_Signal_Delay = float(
        "".join(str(x) for x in re.findall(r"\d+\.\d+", df_inter.iloc[0, 0]))
    )
    Intersection_LOS = "".join(
        str(x) for x in re.findall(r"(?<=:).*", df_inter.iloc[0, 1])
    ).strip()
    print(
        "Intersection_Signal_Delay: ",
        Intersection_Signal_Delay,
        "Intersection_LOS: ",
        Intersection_LOS,
    )

    # Remove last row
    df = df.replace([None], [0], regex=True)
    df = df[:-1]

    # To convert every value to string in the entire DataFrame
    for col in df.columns:
        try:
            df[col] = df[col].astype(str)
        except AttributeError:
            pass

    # get unique values and drop cols with same values in all rows
    nunique = df.nunique()
    cols_to_drop = nunique[nunique == 1].index
    df.drop(cols_to_drop, axis=1, inplace=True)

    #  Transposing for final version
    df = df.T

    df.columns = df.iloc[0]
    df.drop(df.index[0], inplace=True)
    df.columns = df.columns.str.replace(" ", "")
    # convert all column names to lower case
    df.columns = [x.lower() for x in df.columns]
    # if column name doest not exist then set it to blank value
    if "storagelength(ft)" not in df.columns:
        df["storagelength(ft)"] = 0
    if "queuelength95th(ft)" not in df.columns:
        df["queuelength95th(ft)"] = 0
    if "totaldelay" not in df.columns:
        df["totaldelay"] = 0
    if "los" not in df.columns:
        df["los"] = 0
    if "approachdelay" not in df.columns:
        df["approachdelay"] = 0
    if "approachlos" not in df.columns:
        df["approachlos"] = 0

    # remove characters from column
    df["queuelength95th(ft)"] = df["queuelength95th(ft)"].replace(
        "[^0-9]", "", regex=True
    )
    # convert column to numeric and then round of Approach delay and Total delay
    df["totaldelay"] = pd.to_numeric(df["totaldelay"])
    df["approachdelay"] = pd.to_numeric(df["approachdelay"])
    df["storagelength(ft)"] = pd.to_numeric(df["storagelength(ft)"])
    df["queuelength95th(ft)"] = pd.to_numeric(df["queuelength95th(ft)"])
    df["totaldelay"] = df["totaldelay"].apply(lambda x: round(x))
    df["approachdelay"] = df["approachdelay"].apply(lambda x: round(x))

    # replace all 0 from dataframe
    df = df.replace(["0"], [""], regex=True)

    # convert to string
    df["approachlos"] = df["approachlos"].astype(str)
    df["approachdelay"] = df["approachdelay"].astype(str)

    def mergecolumns(x, y):
        if x == "0" or y == "":
            return ""
        else:
            return x + "(" + y + ")"

    # merge ApproachDelay and Approach LOS
    df["approachdelay"] = df.apply(
        lambda x: mergecolumns(x.approachlos, x.approachdelay), axis=1
    )

    # select only required columns from dataframe
    df = df[
        [
            "storagelength(ft)",
            "queuelength95th(ft)",
            "los",
            "totaldelay",
            "approachdelay",
        ]
    ]
    # replace all 0's from below QueueLength95th and LOS with blank
    df["queuelength95th(ft)"] = df["queuelength95th(ft)"].apply(
        lambda x: "" if x == 0 else x
    )
    df["queuelength95th(ft)"] = df["queuelength95th(ft)"].astype(str)
    df["los"] = df["los"].replace([0], [""], regex=True)
    # added IntersectionSignalDelay column to df and set intersection value to the middle cell of DF
    df["intersectionsignaldelay"] = ""
    df["intersectionsignaldelay"][round((len(df.index) - 1) / 2)] = (
        str(Intersection_LOS) + "(" + str(round(Intersection_Signal_Delay)) + ")"
    )

    df["totaldelay"] = df["totaldelay"].apply(lambda x: "--" if x == 0 else x)
    df["storagelength(ft)"] = df["storagelength(ft)"].apply(
        lambda x: "--" if x == 0 else x
    )
    return df


def write_to_word_signalized(list, table, len, shading_elm_1, flag, file_name):
    data = list
    count = 0
    for id, am1, am2, am3, am4, am5, am6, pm1, pm2, pm3, pm4, pm5 in data:
        if id == "0":
            pass
        elif id[:4] in "LaneGroup":
            pass
        else:
            row = table.add_row().cells
            row[0].text = str("")
            row[1].text = str(id)
            if str(am1) == 0:
                row[2].text = ""
            elif str(am1) == "":
                row[2].text = "--"
            else:
                row[2].text = str(am1)
            row[3].text = str(am2)
            row[4].text = str(am3)
            row[5].text = str(am4)
            row[6].text = str(am5)
            row[7].text = str(am6)
            row[8].text = str(pm1)
            row[9].text = str(pm2)
            row[10].text = str(pm3)
            row[11].text = str(pm4)
            row[12].text = str(pm5)
    curr_len = table.rows.__len__()
    arr = []
    for (i, row) in enumerate(table.rows):
        if i > len and i < curr_len:
            tt1 = table.cell(i, 1).text
            tt_prev2 = table.cell(i - 1, 1).text
            tt = table.cell(i, 4).text
            tt_prev = table.cell(i - 1, 4).text
            # if not(tt.strip() and tt_prev.strip()) and tt1[:2] in tt_prev2:
            if (tt.strip() == "--" or tt_prev.strip() == "--") and tt1[:2] in tt_prev2:
                table.cell(
                    i - 1, 1
                ).text = f"{table.cell(i - 1, 1).text}/{table.cell(i, 1).text[-1]}"
                for j in [2, 3, 4, 5, 6, 8, 9, 10, 11]:
                    if (
                        table.cell(i, j).text == "0"
                        or table.cell(i, j).text == "--"
                        or not (table.cell(i, j).text.strip())
                    ):
                        table.cell(i - 1, j).text = table.cell(i - 1, j).text
                    else:
                        table.cell(i - 1, j).text = table.cell(i, j).text
                curr_len = table.rows.__len__()
                arr.append(i)

    # logic for adding custom background color
    for (i, row) in enumerate(table.rows):
        if i > len:
            tt = table.cell(i, 1).text
            tt_prev = table.cell(i - 1, 1).text
            for j in range(0, 12):
                set_table_header_bg_color(table.cell(i, j + 1), shading_elm_1)
                set_table_header_bg_color(table.cell(i - 1, j + 1), shading_elm_1)
            set_table_header_bg_color(table.cell(i - 1, 0), shading_elm_1)

    curr_len_temp = curr_len
    for (i, row) in enumerate(table.rows):
        if i > len and i < curr_len_temp:
            lanegroup_current = table.cell(i, 1).text
            lanegroup_prev = table.cell(i - 1, 1).text
            que_current = table.cell(i, 3).text
            lane_current = table.cell(i, 4).text
            delay_current = table.cell(i, 5).text
            que_prev = table.cell(i - 1, 3).text
            lane_prev = table.cell(i - 1, 4).text
            delay_prev = table.cell(i - 1, 5).text
            if (
                que_current == que_prev
                and delay_current == delay_prev
                and lane_current == lane_prev
            ):
                row = table.rows[i]
                remove_row(table, row)
                curr_len_temp = curr_len_temp - 1

    # for (i,row) in enumerate(table.rows):
    #     if i < arr.__len__():
    #         row = table.rows[arr[i]-i]
    #         #remove_row(table, row)

    for (i, row) in enumerate(table.rows):
        if i > len:
            tt = table.cell(i, 1).text
            tt_prev = table.cell(i - 1, 1).text
            if tt[:2] in tt_prev:
                for j in range(0, 12):
                    a = table.cell(i - 1, j)
                    a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    b = table.cell(i, j)
                    b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    a.merge(b)

    c = table.cell(len, 12)
    d = table.cell(table.rows.__len__() - 1, 12)
    c.merge(d)
    c = table.cell(len, 7)
    d = table.cell(table.rows.__len__() - 1, 7)
    c.merge(d)
    splitted = file_name.split("AM", 1)
    without_underscore = (
        splitted[0].replace("_", " ") if splitted and splitted[0] else ""
    )
    table.cell(len + 2, 0).text = without_underscore
    c = table.cell(len, 0)
    d = table.cell(table.rows.__len__() - 1, 0)
    c.merge(d)

    a = table.cell(0, 2)
    b = table.cell(0, 7)
    a.merge(b)
    a = table.cell(0, 8)
    b = table.cell(0, 12)
    a.merge(b)
    a = table.cell(0, 0)
    b = table.cell(1, 0)
    a.merge(b)
    a = table.cell(0, 1)
    b = table.cell(1, 1)
    a.merge(b)
    return table


def set_table_width(table):
    cell0 = table.columns[0].cells
    cell1 = table.columns[1].cells
    cell2 = table.columns[2].cells
    cell3 = table.columns[3].cells
    cell4 = table.columns[4].cells
    cell5 = table.columns[5].cells
    cell6 = table.columns[6].cells
    cell7 = table.columns[7].cells
    cell8 = table.columns[8].cells
    cell9 = table.columns[9].cells
    cell10 = table.columns[10].cells
    cell11 = table.columns[11].cells
    cell12 = table.columns[12].cells
    cell0.width = Inches(0.8)
    cell1.width = Inches(0.7)
    cell2.width = Inches(0.7)
    cell3.width = Inches(0.65)
    cell4.width = Inches(0.6)
    cell5.width = Inches(0.6)
    cell6.width = Inches(0.75)
    cell7.width = Inches(0.7)
    cell8.width = Inches(0.8)
    cell9.width = Inches(0.6)
    cell10.width = Inches(0.6)
    cell11.width = Inches(0.8)
    cell12.width = Inches(0.8)
    return table


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def set_table_header_bg_color(cell, color_code):
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OxmlElement("w:shd")
    clShading.set(
        qn("w:fill"), color_code
    )  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    tblCellProperties.append(clShading)
    return cell


def read_directory_files_signalized(files, output_file_path):
    print("reading input directory files..")
    even = 0
    flag = True
    AM_file = ""
    df_AM = []
    df_PM = []
    doc = docx.Document()
    paragraph = doc.add_paragraph("Table: Analysis Summary")
    paragraph.alignment = 1
    paragraph.style = "Normal"
    paragraph.runs[0].font.bold = True
    font = paragraph.runs[0].font
    font.size = Pt(11)
    table = doc.add_table(rows=1, cols=13)
    row = table.rows[0].cells
    row[0].text = ""
    row[1].text = ""
    row[2].text = ""
    row[3].text = ""
    row[4].text = ""
    row[5].text = ""
    row[6].text = ""
    row[7].text = ""
    row[8].text = ""
    row[9].text = ""
    row[10].text = ""
    row[11].text = ""
    row[12].text = ""
    row = table.add_row().cells
    row[0].text = "ANALYSIS SCENARIO"
    row[1].text = "LANE GROUP"
    row[2].text = "Existing Storage (ft)"
    row[3].text = "Queue 95th / Max(ft)"
    row[4].text = "Lane LOS"
    row[5].text = "Delay (sec)"
    row[6].text = "Approach LOS (sec)"
    row[7].text = "Overall LOS (sec)"
    row[8].text = "Queue 95th / Max(ft)"
    row[9].text = "Lane LOS"
    row[10].text = "Delay (sec)"
    row[11].text = "Approach LOS (sec)"
    row[12].text = "Overall LOS (sec)"
    table.style = "Table Grid"
    for (i, row) in enumerate(table.rows):
        if i == 1 or i == 2:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        paragraph.runs[0].font.bold = True

        for j in range(0, 12):
            row.cells[7].width = 1200
            row.cells[2].width = 1200

            set_table_header_bg_color(table.cell(i, 0), "EFF0F1")
            set_table_header_bg_color(table.cell(i, j + 1), "EFF0F1")

    row = table.rows[0]
    row.height = Cm(0.7)
    flag = True
    shading_elm_1 = ""
    len = 0

    for (i, file_path) in enumerate(files):
        # size = file_path.seek(0, os.SEEK_END)
        # file_path.seek(0, os.SEEK_SET)
        size = os.path.getsize(file_path)
        if size == 0:
            pass  # if file size is 0/empty then just move on to process next file
        else:
            if i % 2 == 0:
                AM_file = file_path
            if AM_file != file_path:
                if i % 3 == 0:
                    shading_elm_1 = "D4E2EE"
                elif i % 3 == 1:
                    shading_elm_1 = "FFFFFF"
                elif i % 3 == 2:
                    shading_elm_1 = "CDEEBF"
                df_AM = read_text_file_signalized(AM_file)
                df_PM = read_text_file_signalized(file_path)
                file_name = os.path.basename(AM_file)
                # file_name = AM_file.filename
                df_AM["queuelength95th(ft)_pm"] = df_PM["queuelength95th(ft)"]
                df_AM["los_pm"] = df_PM["los"]
                df_AM["totaldelay_pm"] = df_PM["totaldelay"]
                df_AM["approachdelay_pm"] = df_PM["approachdelay"]
                df_AM["intersectionsignaldelay_pm"] = df_PM["intersectionsignaldelay"]
                len = table.rows.__len__()
                file_name = file_name.strip()
                df_AM = df_AM.reset_index()
                df_AM = df_AM.replace(["\(0\)"], [""], regex=True)
                df_AM.rename(columns={2: "LaneGroup"}, inplace=True)
                df_AM = df_AM.replace([""], ["--"], regex=True)
                df_AM["intersectionsignaldelay_pm"] = df_AM[
                    "intersectionsignaldelay_pm"
                ].replace(["--"], [""], regex=True)
                df_AM["intersectionsignaldelay"] = df_AM[
                    "intersectionsignaldelay"
                ].replace(["--"], [""], regex=True)
                df_AM["approachdelay_pm"] = df_AM["approachdelay_pm"].replace(
                    ["--"], [""], regex=True
                )
                df_AM["approachdelay"] = df_AM["approachdelay"].replace(
                    ["--"], [""], regex=True
                )
                table = write_to_word_signalized(
                    df_AM.values.tolist(), table, len, shading_elm_1, flag, file_name
                )

    table.cell(0, 2).text = "Weekday AM Peak Hour"
    table.cell(0, 8).text = "Weekday PM Peak Hour"
    ce1 = table.cell(0, 8)
    ce2 = table.cell(0, 2)
    ce1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ce2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for cell in table.columns[0].cells:
        cell.width = Inches(0.90)
    for cell in table.columns[1].cells:
        cell.width = Inches(0.68)
    for cell in table.columns[2].cells:
        cell.width = Inches(0.66)
    for cell in table.columns[3].cells:
        cell.width = Inches(0.68)
    for cell in table.columns[4].cells:
        cell.width = Inches(0.46)
    for cell in table.columns[5].cells:
        cell.width = Inches(0.50)
    for cell in table.columns[6].cells:
        cell.width = Inches(0.73)
    for cell in table.columns[6].cells:
        cell.width = Inches(0.70)
    for cell in table.columns[7].cells:
        cell.width = Inches(0.62)
    for cell in table.columns[8].cells:
        cell.width = Inches(0.63)
    for cell in table.columns[9].cells:
        cell.width = Inches(0.46)
    for cell in table.columns[10].cells:
        cell.width = Inches(0.50)
    for cell in table.columns[11].cells:
        cell.width = Inches(0.72)
    for cell in table.columns[12].cells:
        cell.width = Inches(0.62)
    for (i, row) in enumerate(table.rows):
        for (j, cell) in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if i < 2:
                        paragraph.runs[0].font.bold = True
                        font = run.font
                        font.size = Pt(9)
                    elif j == 7 or j == 12:
                        font = run.font
                        font.size = Pt(11)
                    else:
                        font = run.font
                        font.size = Pt(9)
                    paragraph.runs[0].font.name = "Times New Roman"
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for paragraph in doc.paragraphs:
        if paragraph.text == "":
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

    now = datetime.now()
    dt_string = now.strftime("%m_%d_%Y_%H:%M:%S")
    file_path = f"{output_file_path}_{dt_string}.docx"
    doc.save(file_path)
    print("completed")
    return file_path


def read_text_file_unsignalized(file):
    df = pd.read_csv(file, sep="delimiter", header=None)

    # merge all columns
    df["combined"] = df.apply(lambda x: "	".join(x.astype(str)), axis=1)

    # drop all other columns
    df = df.drop(columns=df.columns.difference(["combined"]), axis=1)

    # removing nan files from the column
    df["combined"] = df["combined"].str.replace("nan", "0")

    #################################################################

    # Filtered columns
    storage_df = df[
        (
            df["combined"].apply(
                lambda val: any(
                    s in str(val) for s in ["Movement", "Storage Length", "Sign"]
                )
            )
        )
    ]

    # split into columns
    storage_df = storage_df["combined"].str.split("\t", expand=True)

    # Promoted row 1 to header in df
    storage_df.columns = storage_df.iloc[0]
    storage_df.drop(storage_df.index[0], inplace=True)

    # Trim whitespaces in columns and remove empty columns
    storage_df.columns = storage_df.columns.str.strip()
    storage_df = storage_df.loc[:, storage_df.columns != ""]
    storage_df = storage_df.loc[:, storage_df.columns != "0"]
    storage_df = storage_df.loc[:, storage_df.columns != None]

    storage_df = storage_df.T

    # Promoted row Movement to header in storage_df
    storage_df.columns = storage_df.loc["Movement"]
    storage_df.drop(labels=["Movement"], inplace=True)
    storage_df.columns = storage_df.columns.str.strip()

    # if column name doest not exist then set it to blank value
    if "Storage Length" not in storage_df.columns:
        storage_df["Storage Length"] = ""
    if "Sign Control" not in storage_df.columns:
        storage_df["Sign Control"] = ""

    storage_df.rename(
        columns={"Storage Length": "storagelength(ft)", "Sign Control": "signcontrol"},
        inplace=True,
    )

    storage_df["storagelength(ft)"].replace(to_replace="0", value="-", inplace=True)

    storage_df["Movement"] = storage_df.index

    # Reset Index
    storage_df.reset_index(drop=True, inplace=True)

    #################################################################

    # get approach values
    start_index = df.index[df["combined"].str.contains("Approach")][0]
    stop_index = start_index + 3

    approach_df = df[start_index:stop_index]

    # split into columns
    approach_df = approach_df["combined"].str.split("\t", expand=True)

    # get unique values and drop cols with same values in all rows
    nunique = approach_df.nunique()
    cols_to_drop = nunique[nunique <= 2].index
    approach_df.drop(cols_to_drop, axis=1, inplace=True)

    # replace all null cols from dataframe
    approach_df = approach_df.replace([""], [None], regex=True)
    approach_df.dropna(axis=1, inplace=True)

    # Promoted row 1 to header in df
    approach_df.columns = approach_df.iloc[0]
    approach_df.drop(approach_df.index[0], inplace=True)
    # Trim whitespaces in columns and remove empty columns
    approach_df.columns = approach_df.columns.str.strip()

    approach_df.columns = (
        approach_df.columns.str.replace("SB", "SBL")
        .str.replace("NB", "NBL")
        .str.replace("EB", "EBL")
    )

    approach_df = approach_df.T

    # Promoted row 1 to header in df
    approach_df.columns = approach_df.iloc[0]
    approach_df.drop(approach_df.index[0], inplace=True)
    # Trim whitespaces in columns and remove empty columns
    approach_df.columns = approach_df.columns.str.strip()

    approach_df.rename(
        columns={"HCM Control Delay, s": "approachdelay", "HCM LOS": "approachlos"},
        inplace=True,
    )

    if "approachdelay" not in approach_df.columns:
        approach_df["approachdelay"] = 0
    if "approachlos" not in approach_df.columns:
        approach_df["approachlos"] = 0

    # remove special symbols
    approach_df["approachdelay"] = approach_df["approachdelay"].apply(
        lambda x: re.sub(r"[$,+*]", "", x)
    )
    try:
        # convert column to numeric and then round of Approach delay and Total delay
        approach_df["approachdelay"] = pd.to_numeric(approach_df["approachdelay"])
        approach_df["approachdelay"] = approach_df["approachdelay"].apply(
            lambda x: round(x)
        )
        # convert to string
        approach_df["approachlos"] = approach_df["approachlos"].astype(str)
        approach_df["approachdelay"] = approach_df["approachdelay"].astype(str)
    except:
        print("Error printing numeric")

    def mergecolumns(x, y):
        if x == "0" or y == "":
            return ""
        else:
            return x + " (" + y + ")"

    # merge ApproachDelay and Approach LOS
    approach_df["approachdelay"] = approach_df.apply(
        lambda x: mergecolumns(x.approachlos, x.approachdelay), axis=1
    )

    approach_df["Movement"] = approach_df.index

    # Reset Index
    approach_df.reset_index(drop=True, inplace=True)

    #################################################################

    # get HCM values

    start_index = df.index[df["combined"].str.contains("Minor")][1]  # second occurence

    HCM_df = df[36:]

    # Filtered combined
    HCM_df = HCM_df[
        (
            HCM_df["combined"].apply(
                lambda val: any(s in str(val) for s in ["LOS", "Minor", "Delay", "95"])
            )
        )
    ]

    # split into columns
    HCM_df = HCM_df["combined"].str.split("\t", expand=True)

    # get unique values and drop cols with same values in all rows
    nunique = HCM_df.nunique()
    cols_to_drop = nunique[nunique <= 2].index
    HCM_df.drop(cols_to_drop, axis=1, inplace=True)

    # Promoted row 1 to header in df
    HCM_df.columns = HCM_df.iloc[0]
    HCM_df.drop(HCM_df.index[0], inplace=True)

    HCM_df.columns = (
        HCM_df.columns.str.replace("SBLn1", "SBL")
        .str.replace("SBLn2", "SBR")
        .str.replace("NBLn1", "NBL")
        .str.replace("NBLn2", "NBR")
        .str.replace("EBLn1", "EBL")
        .str.replace("EBLn2", "EBR")
    )

    HCM_df = HCM_df.T

    # Promoted row 1 to header in df
    HCM_df.columns = HCM_df.iloc[0]
    HCM_df.drop(HCM_df.index[0], inplace=True)
    # Trim whitespaces in columns and remove empty columns
    HCM_df.columns = HCM_df.columns.str.strip()

    # Error handling if columns are missing
    if "HCM Control Delay (s)" not in HCM_df.columns:
        HCM_df["HCM Control Delay (s)"] = ""
    if "HCM Lane LOS" not in HCM_df.columns:
        HCM_df["HCM Lane LOS"] = ""
    if "HCM 95th %tile Q(veh)" not in HCM_df.columns:
        HCM_df["HCM 95th %tile Q(veh)"] = ""

    HCM_df.rename(
        columns={
            "HCM Control Delay (s)": "totaldelay",
            "HCM Lane LOS": "los",
            "HCM 95th %tile Q(veh)": "queuelength95th(ft)",
        },
        inplace=True,
    )
    # remove special symbols
    HCM_df["totaldelay"] = HCM_df["totaldelay"].apply(
        lambda x: re.sub(r"[$,+*]", "", x)
    )
    try:
        import math

        # convert column to numeric and then round of Approach delay and Total delay
        HCM_df["queuelength95th(ft)"] = pd.to_numeric(HCM_df["queuelength95th(ft)"])
        HCM_df["queuelength95th(ft)"] = HCM_df["queuelength95th(ft)"].apply(
            lambda x: math.ceil(x * 25)
        )
        HCM_df["totaldelay"] = pd.to_numeric(HCM_df["totaldelay"])
        HCM_df["totaldelay"] = HCM_df["totaldelay"].apply(lambda x: round(x))
    except:
        print("Error in converting to numeric")

    # convert to string
    HCM_df["queuelength95th(ft)"] = HCM_df["queuelength95th(ft)"].astype(str)
    HCM_df["totaldelay"] = HCM_df["totaldelay"].astype(str)
    HCM_df["Movement"] = HCM_df.index

    # Reset Index
    HCM_df.reset_index(drop=True, inplace=True)

    #################################################################

    # Merged storage_df and HCM_df into merged_st_hcm
    merged_st_hcm = storage_df.merge(
        HCM_df, left_on=["Movement"], right_on=["Movement"], how="outer"
    )
    # Merged merged_st_hcm and approach_df into merged_st_hcm_ap
    merged_st_hcm_ap = merged_st_hcm.merge(
        approach_df,
        left_on=["Movement"],
        right_on=["Movement"],
        how="outer",
        suffixes=["_storage_df", "_HCM_df"],
    )

    # create final df
    final_df = merged_st_hcm_ap

    # Add overall column
    if "overalllos" not in final_df.columns:
        final_df["overalllos"] = ""

    # Replace Movement to Lanegroup
    final_df.rename(columns={"Movement": "lanegroup"}, inplace=True)

    # create final dataframe
    final_df = final_df[
        [
            "lanegroup",
            "storagelength(ft)",
            "queuelength95th(ft)",
            "los",
            "totaldelay",
            "approachdelay",
            "overalllos",
        ]
    ]

    # Replace None with -
    final_df = final_df.replace([None], ["-"], regex=True)
    final_df.fillna("-", inplace=True)
    # Remove rows with '-' in langroup
    final_df = final_df[final_df["lanegroup"] != "-"]
    # Remove rows where all columns have '-'
    final_df = final_df[~final_df.iloc[:, 1:6].eq("-").all(1)]
    # Reset Index
    final_df.reset_index(drop=True, inplace=True)
    # Add N/A to the overall los
    final_df["overalllos"][round((len(final_df.index) - 1) / 2)] = "N/A"

    return final_df


def myfunc(x, y):
    if x == "-" or x == "":
        return x
    elif y == "Free":
        return f"{x}\N{SUPERSCRIPT ONE}"
    else:
        return f"{x}\N{SUPERSCRIPT TWO}"


def myfuncApproachLOS(x, y, z):
    if x == "-" or x == "--" or x == "":
        return x
    elif z == "Free":
        return f"{y}({x})\N{SUPERSCRIPT ONE}"
    else:
        return f"{y}({x})\N{SUPERSCRIPT TWO}"


def write_to_word_unsignalized(list, table, len, shading_elm_1, flag, file_name):
    data = list
    count = 0
    for id, am1, am2, am3, am4, am5, am6, pm1, pm2, pm3, pm4, pm5 in data:
        if id == "0":
            pass
        elif id[:4] in "lanegroup":
            pass
        else:
            row = table.add_row().cells
            row[0].text = str("")
            row[1].text = str(id)
            if str(am1) == 0:
                row[2].text = ""
            elif str(am1) == "":
                row[2].text = "--"
            else:
                row[2].text = str(am1)
            row[3].text = str(am2)
            row[4].text = str(am3)
            row[5].text = str(am4)
            row[6].text = str(am5)
            row[7].text = str(am6)
            row[8].text = str(pm1)
            row[9].text = str(pm2)
            row[10].text = str(pm3)
            row[11].text = str(pm4)
            row[12].text = str(pm5)
    curr_len = table.rows.__len__()
    arr = []
    for (i, row) in enumerate(table.rows):
        if i > len and i < curr_len:
            tt1 = table.cell(i, 1).text
            tt_prev2 = table.cell(i - 1, 1).text
            tt = table.cell(i, 4).text
            tt_prev = table.cell(i - 1, 4).text
            # if not(tt.strip() and tt_prev.strip()) and tt1[:2] in tt_prev2:
            if (tt.strip() == "--" or tt_prev.strip() == "--") and tt1[:2] in tt_prev2:
                table.cell(
                    i - 1, 1
                ).text = f"{table.cell(i - 1, 1).text}/{table.cell(i, 1).text[-1]}"
                for j in [2, 3, 4, 5, 6, 8, 9, 10, 11]:
                    if table.cell(i, j).text == "0" or table.cell(i, j).text == "--":
                        table.cell(i - 1, j).text = table.cell(i - 1, j).text
                    else:
                        table.cell(i - 1, j).text = table.cell(i, j).text
                curr_len = table.rows.__len__()
                arr.append(i)

    # logic for adding custom background color
    for (i, row) in enumerate(table.rows):
        if i > len:
            tt = table.cell(i, 1).text
            tt_prev = table.cell(i - 1, 1).text
            for j in range(0, 12):
                set_table_header_bg_color(table.cell(i, j + 1), shading_elm_1)
                set_table_header_bg_color(table.cell(i - 1, j + 1), shading_elm_1)
            set_table_header_bg_color(table.cell(i - 1, 0), shading_elm_1)

    curr_len_temp = curr_len
    for (i, row) in enumerate(table.rows):
        if i > len and i < curr_len_temp:
            lanegroup_current = table.cell(i, 1).text
            lanegroup_prev = table.cell(i - 1, 1).text
            que_current = table.cell(i, 3).text
            lane_current = table.cell(i, 4).text
            delay_current = table.cell(i, 5).text
            que_prev = table.cell(i - 1, 3).text
            lane_prev = table.cell(i - 1, 4).text
            delay_prev = table.cell(i - 1, 5).text
            if (
                que_current == que_prev
                and delay_current == delay_prev
                and lane_current == lane_prev
            ) or table.cell(i, 1).text == "":
                row = table.rows[i]
                remove_row(table, row)
                curr_len_temp = curr_len_temp - 1

    for (i, row) in enumerate(table.rows):
        if i > len:
            tt = table.cell(i, 1).text
            tt_prev = table.cell(i - 1, 1).text
            if tt[:2] in tt_prev:
                for j in range(0, 12):
                    a = table.cell(i - 1, j)
                    a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    b = table.cell(i, j)
                    b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    a.merge(b)

    c = table.cell(len, 12)
    d = table.cell(table.rows.__len__() - 1, 12)
    c.merge(d)
    c = table.cell(len, 7)
    d = table.cell(table.rows.__len__() - 1, 7)
    c.merge(d)
    splitted = file_name.split("AM", 1)
    without_underscore = (
        splitted[0].replace("_", " ") if splitted and splitted[0] else ""
    )
    table.cell(len + 1, 0).text = without_underscore
    c = table.cell(len, 0)
    d = table.cell(table.rows.__len__() - 1, 0)
    c.merge(d)

    a = table.cell(0, 3)
    b = table.cell(0, 7)
    a.merge(b)
    a = table.cell(0, 8)
    b = table.cell(0, 12)
    a.merge(b)
    a = table.cell(0, 0)
    b = table.cell(1, 0)
    a.merge(b)
    a = table.cell(0, 1)
    b = table.cell(1, 1)
    a.merge(b)
    a = table.cell(0, 2)
    b = table.cell(1, 2)
    a.merge(b)
    return table


def read_directory_files_unsignalized(files, output_file_path):
    print("reading input directory files..")
    even = 0
    flag = True
    AM_file = ""
    df_AM = []
    df_PM = []
    doc = docx.Document()
    paragraph = doc.add_paragraph("Table: Analysis Summary")
    paragraph.alignment = 1
    paragraph.style = "Normal"
    paragraph.runs[0].font.bold = True
    font = paragraph.runs[0].font
    font.size = Pt(11)
    table = doc.add_table(rows=1, cols=13)
    row = table.rows[0].cells
    row[0].text = ""
    row[1].text = ""
    row[2].text = ""
    row[3].text = ""
    row[4].text = ""
    row[5].text = ""
    row[6].text = ""
    row[7].text = ""
    row[8].text = ""
    row[9].text = ""
    row[10].text = ""
    row[11].text = ""
    row[12].text = ""
    row = table.add_row().cells
    row[0].text = "ANALYSIS SCENARIO"
    row[1].text = "LANE GROUP"
    row[2].text = "Existing Storage (ft)"
    row[3].text = "Queue 95th / Max(ft)"
    row[4].text = "Lane LOS"
    row[5].text = "Delay (sec)"
    row[6].text = "Approach LOS (sec)"
    row[7].text = "Overall LOS (sec)"
    row[8].text = "Queue 95th / Max(ft)"
    row[9].text = "Lane LOS"
    row[10].text = "Delay (sec)"
    row[11].text = "Approach LOS (sec)"
    row[12].text = "Overall LOS (sec)"
    table.style = "Table Grid"
    for (i, row) in enumerate(table.rows):
        if i == 1 or i == 2:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        paragraph.runs[0].font.bold = True
        for j in range(0, 12):
            row.cells[7].width = 1200
            row.cells[2].width = 1200

            set_table_header_bg_color(table.cell(i, 0), "EFF0F1")
            set_table_header_bg_color(table.cell(i, j + 1), "EFF0F1")
    row = table.rows[0]
    row.height = Cm(0.7)
    flag = True
    shading_elm_1 = ""
    len = 0

    for (i, file_path) in enumerate(files):
        size = file_path.seek(0, os.SEEK_END)
        file_path.seek(0, os.SEEK_SET)
        #size = os.path.getsize(file_path)
        if size == 0:
            pass  # if file size is 0/empty then just move on to process next file
        else:
            if i % 2 == 0:
                AM_file = file_path
            if AM_file != file_path:
                if i % 3 == 0:
                    shading_elm_1 = "D4E2EE"
                elif i % 3 == 1:
                    shading_elm_1 = "FFFFFF"
                elif i % 3 == 2:
                    shading_elm_1 = "CDEEBF"
                df_AM = read_text_file_unsignalized(AM_file)
                df_PM = read_text_file_unsignalized(file_path)
                #file_name = os.path.basename(AM_file)
                file_name = AM_file.filename
                df_AM["queuelength95th(ft)_pm"] = df_PM["queuelength95th(ft)"]
                df_AM["los_pm"] = df_PM["los"]
                df_AM["totaldelay_pm"] = df_PM["totaldelay"]
                df_AM["approachdelay_pm"] = df_PM["approachdelay"]
                df_AM["overalllos_pm"] = df_PM["overalllos"]
                len = table.rows.__len__()
                file_name = file_name.strip()
                # df_AM = df_AM.reset_index()
                df_AM = df_AM.replace(["\(0\)"], [""], regex=True)
                df_AM = df_AM.replace([""], ["--"], regex=True)
                df_AM["queuelength95th(ft)_pm"] = df_AM[
                    "queuelength95th(ft)_pm"
                ].replace(["--"], [""], regex=True)
                df_AM["overalllos"] = df_AM["overalllos"].replace(
                    ["--"], [""], regex=True
                )
                df_AM["overalllos_pm"] = df_AM["overalllos_pm"].replace(
                    ["--"], [""], regex=True
                )
                df_AM["approachdelay_pm"] = df_AM["approachdelay_pm"].replace(
                    ["--"], [""], regex=True
                )
                df_AM["approachdelay"] = df_AM["approachdelay"].replace(
                    ["--"], [""], regex=True
                )
                table = write_to_word_unsignalized(
                    df_AM.values.tolist(), table, len, shading_elm_1, flag, file_name
                )

    table.cell(0, 3).text = "Weekday AM Peak Hour11"
    table.cell(0, 9).text = "Weekday PM Peak Hour"
    ce1 = table.cell(0, 9)
    ce2 = table.cell(0, 3)
    ce1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ce2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.alignment = WD_ALIGN_VERTICAL.CENTER
    table.autofit = False
    table.allow_autofit = False
    for cell in table.columns[0].cells:
        cell.width = Inches(0.90)
    for cell in table.columns[1].cells:
        cell.width = Inches(0.68)
    for cell in table.columns[2].cells:
        cell.width = Inches(0.66)
    for cell in table.columns[3].cells:
        cell.width = Inches(0.68)
    for cell in table.columns[4].cells:
        cell.width = Inches(0.46)
    for cell in table.columns[5].cells:
        cell.width = Inches(0.50)
    for cell in table.columns[6].cells:
        cell.width = Inches(0.73)
    for cell in table.columns[6].cells:
        cell.width = Inches(0.70)
    for cell in table.columns[7].cells:
        cell.width = Inches(0.62)
    for cell in table.columns[8].cells:
        cell.width = Inches(0.63)
    for cell in table.columns[9].cells:
        cell.width = Inches(0.46)
    for cell in table.columns[10].cells:
        cell.width = Inches(0.50)
    for cell in table.columns[11].cells:
        cell.width = Inches(0.72)
    for cell in table.columns[12].cells:
        cell.width = Inches(0.62)
    for (i, row) in enumerate(table.rows):
        for (j, cell) in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if i < 2:
                        paragraph.runs[0].font.bold = True
                        font = run.font
                        font.size = Pt(9)
                    elif j == 7 or j == 12:
                        font = run.font
                        font.size = Pt(11)
                    else:
                        font = run.font
                        font.size = Pt(9)
                    paragraph.runs[0].font.name = "Times New Roman"
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    len = table.rows.__len__()
    # for remove extra \n from analysis scenario column
    for (i, row) in enumerate(table.rows):
        if i > 2:
            temp1 = table.columns[0].cells[i].text
            temp2 = table.columns[7].cells[i].text
            temp3 = table.columns[12].cells[i].text
            # temp12= table.columns[11].cells[i].text
            table.columns[0].cells[i].text = temp1.replace("\n", "")
            if temp2[-2] == '\n':
                table.columns[7].cells[i].text = temp2[:5]
            if temp3[-2] == '\n':
                table.columns[12].cells[i].text = temp3[:5]
            # table.columns[11].cells[i].text= (temp12.replace("\n", ""))

    for paragraph in doc.paragraphs:
        if paragraph.text == "":
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

    doc.add_paragraph(
        "\n*Due to limitations with HCM 6th TWSC reporting, a vehicle length of 25 feet was used to determine the 95th percentile queue in feet.\n1.	Level of service for major street left-turn movement.\n2.	Level of service for minor-street approach."
    )
    doc.add_heading("NOTE:", 3)
    doc.add_paragraph(
        "95th Queue in Table is calculated as follows:  25 x HMC 95th %tile Q (veh)"
    )

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(9)

    now = datetime.now()
    #dt_string = now.strftime("%m_%d_%Y_%H:%M:%S")
    dt_string = now.strftime("%d-%m-%Y_%H%M")
    file_path = f"{output_file_path}_{dt_string}.docx"
    doc.save(file_path)
    print("completed")
    return file_path
